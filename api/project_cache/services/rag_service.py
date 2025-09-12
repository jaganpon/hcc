import os
import json
import hashlib
import re
from pathlib import Path
from typing import List, Optional
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from sqlalchemy.orm import Session
from dotenv import load_dotenv
import logging

from api.project_cache.models import Document, Chunk, FAQ

# ---------------- Load Config ----------------
load_dotenv()

STORAGE_DIR = Path(os.getenv("STORAGE_DIR", "./storage"))
STORAGE_DIR.mkdir(parents=True, exist_ok=True)

PROVIDER = os.getenv("OPENAI_PROVIDER", "openai").lower()

# ---------------- Clients ----------------
if PROVIDER == "azure":
    from openai import AzureOpenAI
    client = AzureOpenAI(
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    )
    CHAT_MODEL = os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT")
    EMBED_MODEL = os.getenv("AZURE_OPENAI_EMBED_DEPLOYMENT")
    logging.info("🔹 Using **Azure OpenAI** as the LLM provider")

else:  # default → standard OpenAI
    from openai import OpenAI
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    CHAT_MODEL = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")
    EMBED_MODEL = os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")
    logging.info("🔹 Using **OpenAI (non-Azure)** as the LLM provider")

# ---------------- Utility ----------------
def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()

def chunk_text(text: str, max_chars: int = 1000, overlap: int = 150) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = end - overlap
    return chunks

def load_text_from_file(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suf == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    if suf == ".docx":
        doc = DocxDocument(str(path))
        return "\n".join([p.text for p in doc.paragraphs])
    raise ValueError("Unsupported file type")

def embed_texts(texts: list[str]) -> list[list[float]]:
    if not texts:
        return []
    resp = client.embeddings.create(model=EMBED_MODEL, input=texts)
    return [item.embedding for item in resp.data]

def cosine(a: list[float], b: list[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = (sum(x * x for x in a) ** 0.5) or 1e-9
    nb = (sum(y * y for y in b) ** 0.5) or 1e-9
    return dot / (na * nb)

# ---------------- Storage / DB ----------------
def save_file_and_chunks(db: Session, user_id: Optional[str], file_bytes: bytes, filename: str) -> int:
    user_dir = STORAGE_DIR / (user_id or "shared")
    user_dir.mkdir(parents=True, exist_ok=True)

    base, ext = os.path.splitext(filename)
    target = user_dir / filename
    i = 1
    while target.exists():
        target = user_dir / f"{base}({i}){ext}"
        i += 1

    with open(target, "wb") as f:
        f.write(file_bytes)

    fhash = sha256_file(target)

    doc = Document(user_id=user_id, filename=target.name, filepath=str(target), filehash=fhash)
    db.add(doc)
    db.commit()
    db.refresh(doc)

    text = load_text_from_file(target)
    parts = chunk_text(text)
    embs = embed_texts(parts)

    for ord_i, (t, e) in enumerate(zip(parts, embs)):
        ch = Chunk(document_id=doc.id, user_id=user_id, text=t, embedding=json.dumps(e), ord=ord_i)
        db.add(ch)
    db.commit()
    return doc.id

def list_files(db: Session, user_id: Optional[str]):
    q = db.query(Document)
    if user_id:
        q = q.filter(Document.user_id == user_id)
    rows = q.order_by(Document.created_at.desc()).all()
    return [{"id": d.id, "filename": d.filename, "created_at": d.created_at} for d in rows]

def clear_files(db: Session, user_id: Optional[str]):
    docs = db.query(Document).filter(Document.user_id == user_id if user_id else True).all()
    for d in docs:
        try:
            Path(d.filepath).unlink(missing_ok=True)
        except Exception:
            pass
        db.delete(d)
    db.commit()

def upsert_faq(db: Session, user_id: Optional[str], question: str):
    qnorm = question.strip().lower()
    existing = db.query(FAQ).filter(FAQ.question == qnorm).first()
    if existing:
        existing.count += 1
    else:
        existing = FAQ(user_id=user_id, question=qnorm, count=1)
        db.add(existing)
    db.commit()

def search(db: Session, user_id: Optional[str], query: str, k: int = 5):
    qvecs = embed_texts([query])
    if not qvecs:
        return []
    qvec = qvecs[0]
    chunks = db.query(Chunk).filter(Chunk.user_id == user_id if user_id else True).all()
    scored = []
    import json as _json
    for c in chunks:
        try:
            emb = _json.loads(c.embedding)
            scored.append((c, cosine(emb, qvec)))
        except Exception:
            continue
    scored.sort(key=lambda x: x[1], reverse=True)
    return scored[:k]

# ---------------- Chat ----------------
def build_answer(context_snippets: list[str], question: str) -> str:
    system = (
        "You are HR onboarding assistant. Answer using ONLY the provided context. "
        "If the answer is not in the context, say you have limited information and suggest contacting HR for more info."
    )
    prompt = "Context:\n" + "\n---\n".join(context_snippets) + f"\n\nQuestion: {question}"
    chat = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    return chat.choices[0].message.content.strip()

DEFAULT_USER = "hr_admin"

def handle_chat(db: Session, user_id: Optional[str], session_id: str, message: str):
    try:
        if message and len(message.split()) >= 2:
            upsert_faq(db, user_id or DEFAULT_USER, message)

        top = search(db, DEFAULT_USER, message, k=5)

        snippets = []
        sources = []

        if top:
            for chunk, score in top:
                snippets.append(chunk.text[:1200])
                doc = db.get(Document, chunk.document_id)
                sources.append({
                    "document_id": chunk.document_id,
                    "filename": doc.filename if doc else "unknown",
                    "snippet": chunk.text[:200]
                })
            answer = build_answer(snippets, message)
            return answer, sources
        else:
            response = client.chat.completions.create(
                model=CHAT_MODEL,
                messages=[
                    {"role": "system", "content": "You are a helpful HR assistant."},
                    {"role": "user", "content": message},
                ],
                temperature=0.7,
                max_tokens=500,
            )
            answer = response.choices[0].message.content.strip()
            return answer, []

    except Exception as e:
        print(f"[handle_chat ERROR] {e}")
        return f"Sorry, something went wrong: {str(e)}", []
